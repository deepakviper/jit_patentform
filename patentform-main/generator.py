import os
import re
import json
# pyrefly: ignore [missing-import]
import docx

class FormGenerator:
    """Generates a structured .docx Form 1 file by replacing placeholders in a template."""
    
    def __init__(self, template_path: str, static_data_path: str):
        self.template_path = template_path
        
        # Load the JSON static replacements
        if not os.path.exists(static_data_path):
            raise FileNotFoundError(f"Static config file not found: {static_data_path}")
            
        with open(static_data_path, 'r', encoding='utf-8') as f:
            try:
                self.static_replacements = json.load(f)
            except json.JSONDecodeError as e:
                raise ValueError(f"Error reading 'static_data.json'. The file is likely empty or invalid. Please check its contents. (Details: {e})")

    def generate_docx(self, metadata: dict, output_path: str):
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template file '{self.template_path}' not found.")

        doc = docx.Document(self.template_path)

        merged_data = {**self.static_replacements, **metadata}
        
        # Safely extract nested data for flattening with default fallbacks
        applicant = merged_data.get('applicant') or {}
        if not isinstance(applicant, dict):
            applicant = {}
        app_address = applicant.get('address') or {}
        if not isinstance(app_address, dict):
            app_address = {}
            
        inventors = merged_data.get('inventors') or []
        if not isinstance(inventors, list):
            inventors = []
            
        inv_address = merged_data.get('inventor_address') or {}
        if not isinstance(inv_address, dict):
            inv_address = {}
            
        service = merged_data.get('address_for_service') or {}
        if not isinstance(service, dict):
            service = {}

        # Safely format list variables to prevent join() of None types
        inv_names = []
        inv_nationalities = []
        inv_countries = []
        for inv in inventors:
            if isinstance(inv, dict):
                inv_names.append(str(inv.get('name') or ''))
                inv_nationalities.append(str(inv.get('nationality') or 'Indian'))
                inv_countries.append(str(inv.get('country_of_residence') or 'India'))

        replacements = {
            "{{APP_NAME}}": applicant.get('name') or '',
            "{{NATIONALITY}}": applicant.get('nationality') or '',
            "{{RES_CO}}": applicant.get('country_of_residence') or '',
            "{{HOUSE_NO}}": app_address.get('house_no') or '',
            "{{STREET}}": app_address.get('street') or '',
            "{{CITY}}": app_address.get('city') or '',
            "{{STATE}}": app_address.get('state') or '',
            "{{COUNTRY}}": app_address.get('country') or '',
            "{{PINCODE}}": app_address.get('pin_code') or '',
            
            "{{TITLE}}": (merged_data.get('title_of_invention') or '').upper(),
            
            "{{INV_NAME}}": "\n".join(inv_names),
            "{{INV_NAT}}": "\n".join(inv_nationalities),
            "{{INV_COUNTRY}}": "\n".join(inv_countries),
            
            "{{INV_HOUSE_NO}}": inv_address.get('house_no') or '',
            "{{INV_STREET}}": inv_address.get('street') or '',
            "{{INV_CITY}}": inv_address.get('city') or '',
            "{{INV_STATE}}": inv_address.get('state') or '',
            "{{INV_COUNTRY_ADDR}}": inv_address.get('country') or '',
            "{{INV_PIN}}": inv_address.get('pin_code') or '',
            
            "{{SERVICE_NAME}}": service.get('name') or '',
            "{{SERVICE_ADDRESS}}": service.get('postal_address') or '',
            "{{SERVICE_TEL}}": service.get('telephone_no') or '',
            "{{SERVICE_MOBILE}}": service.get('mobile_no') or '',
            "{{SERVICE_FAX}}": service.get('fax_no') or '',
            "{{SERVICE_EMAIL}}": service.get('email_id') or ''
        }

        # --- NEW FORM 2 MAPPINGS ---
        # 1. Format the full description dynamically from the extracted sections
        desc_sections = []
        desc_data = merged_data.get('description') or {}
        if isinstance(desc_data, str):
            desc_sections.append(desc_data)
        elif isinstance(desc_data, dict):
            for key, val in desc_data.items():
                if val:
                    # Capitalize and format headers nicely
                    header = key.replace('_', ' ').upper()
                    if header == "TECHNICAL FIELD": header = "TECHNICAL FIELD OF THE INVENTION"
                    if header == "BACKGROUND": header = "BACKGROUND OF THE INVENTION"
                    if header == "OBJECTIVES": header = "OBJECTIVES OF THE INVENTION"
                    desc_sections.append(f"{header}\n{val}")
        
        replacements["{{DESCRIPTION}}"] = "\n\n".join(desc_sections)
        
        # 2. Add Abstract
        replacements["{{ABSTRACT}}"] = merged_data.get('abstract') or ''
        
        # 3. Create a combined single-string address for the inventor
        address_parts = [
            inv_address.get('house_no') or '',
            inv_address.get('street') or '',
            inv_address.get('city') or '',
            inv_address.get('state') or '',
            inv_address.get('country') or '',
            inv_address.get('pin_code') or ''
        ]
        # Filter out empty or literal "None" strings and join with commas
        valid_parts = [str(p).strip() for p in address_parts if p and str(p).lower() != 'none']
        replacements["{{INV_FULL_ADDRESS}}"] = ", ".join(valid_parts).strip()
        
        # 4. Final safety sweep: replace any remaining None values with empty strings
        for key, val in replacements.items():
            if val is None or str(val).lower() == 'none':
                replacements[key] = ""
        # ----------------------------

        def replace_in_paragraph(p):
            """Helper function to find and replace placeholders in a single paragraph."""
            if '{{' in p.text:
                text = p.text
                
                # Replace known data mapping
                for key, val in replacements.items():
                    text = text.replace(key, str(val))
                
                # Clear out any leftover placeholders that had no data provided
                # This ensures fields are left completely blank as requested
                text = re.sub(r'\{\{.*?\}\}', '', text)
                
                if p.text != text:
                    p.text = text

        # Replace in all document paragraphs (standard text)
        for p in doc.paragraphs:
            replace_in_paragraph(p)

        # Replace in all table cells (which is where most Form 1 data lives)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

        # Save the finalized document
        doc.save(output_path)